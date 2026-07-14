from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "/Users/surachartlimrattanaphun/Desktop/tour/docs/Booking_Summary_A4_Spec.docx"


COLORS = {
    "navy": "1F3448",
    "teal": "18B89B",
    "light_teal": "E8F7F4",
    "light_blue": "EAF3FB",
    "light_gray": "F3F5F7",
    "border": "D9E1E8",
    "green": "1F8F46",
    "muted": "5D6B78",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=COLORS["border"], size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="000000", size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def style_table(table, header=True):
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
        if header and row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, COLORS["navy"])
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        size = 15
        color = COLORS["navy"]
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        size = 11
        color = COLORS["navy"]
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_body(doc, text, size=9.5, color="000000", after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("• ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COLORS["teal"])
    body = p.add_run(text)
    body.font.name = "Arial"
    body.font.size = Pt(9)
    return p


def add_key_value_table(doc, rows, widths=(4.2, 5.0, 4.2, 5.0)):
    table = doc.add_table(rows=len(rows), cols=4)
    table.autofit = False
    set_col_widths(table, list(widths))
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            is_label = c in (0, 2)
            set_cell_text(
                table.cell(r, c),
                value,
                bold=is_label,
                color=COLORS["muted"] if is_label else "000000",
                size=8.5,
            )
            set_cell_shading(table.cell(r, c), COLORS["light_gray"] if is_label else "FFFFFF")
    style_table(table, header=False)
    return table


def add_service_table(doc, title, headers, rows, widths):
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=len(headers))
    set_col_widths(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF", size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), COLORS["navy"])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER if i not in (2, 3) else WD_ALIGN_PARAGRAPH.LEFT)
            if row[-1] == "Confirmed":
                set_cell_shading(cells[i], "EEF9F0")
    style_table(table, header=False)
    return table


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Booking Confirmed Summary")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("A4 printout requirement and document mockup for Admin Booking")
    sr.font.name = "Arial"
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor.from_string(COLORS["muted"])

    add_heading(doc, "1. Requested Feature", 1)
    add_body(
        doc,
        "Add a Summary button at the bottom of the Admin Edit Booking page, near the Booking Confirmed action. "
        "When the booking is confirmed, admin should be able to generate a clean A4 summary of the confirmed booking and all booked services.",
    )
    add_bullet(doc, "The Summary document must be printable on A4 portrait.")
    add_bullet(doc, "The printout must show booking information only; it must not show edit, delete, approve, email, or UI action buttons.")
    add_bullet(doc, "The summary should include all service sections used in the booking: transfers, hotels, excursions, tours, flights, special packages, and others when present.")
    add_bullet(doc, "Confirmed services should be visually clear, using a restrained green status style.")

    add_heading(doc, "2. Button Placement and Behavior", 1)
    rows = [
        ["Page", "Admin > Edit Booking", "Button Label", "Summary"],
        ["Position", "Bottom action area, beside Booking Confirmed", "Visibility", "Admin only"],
        ["Enabled When", "Booking status is Confirmed or all service rows are confirmed", "Output", "A4 printable summary"],
        ["Recommended Action", "Open print/PDF preview or generate PDF/DOCX", "Language", "English"],
    ]
    add_key_value_table(doc, rows)

    add_heading(doc, "3. Data Required in the A4 Summary", 1)
    requirement_table = doc.add_table(rows=1, cols=3)
    set_col_widths(requirement_table, [4.3, 8.8, 5.1])
    for i, h in enumerate(["Section", "Fields", "Notes"]):
        set_cell_text(requirement_table.cell(0, i), h, bold=True, color="FFFFFF", size=8.3, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(requirement_table.cell(0, i), COLORS["navy"])
    req_rows = [
        ("Booking Header", "Booking ref, quotation ref, agent, client name, booking date, trip start/end, pax, phone, status", "Always required"),
        ("Transfer Details", "Date/flight, transfer type, city, description, TOT, from, to, remarks, status", "Only show rows included in booking"),
        ("Hotel Details", "Check-in, check-out, city, hotel, room type, nights, single room, double room, extra bed, promotion, status", "Include special package badge if relevant"),
        ("Excursion Details", "Date, city, name, pickup time, hotel, remarks, TOE, status", "Confirmed rows should show status"),
        ("Tour Details", "Start date, end date, city, name, TOT, route, pax, remarks, price/status", "Include tour accommodation notes if present"),
        ("Internal Notes", "Booking remarks and notes for office/admin", "Optional; can be hidden in customer version"),
    ]
    for row in req_rows:
        cells = requirement_table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.2)
    style_table(requirement_table, header=False)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A4 Booking Summary Mockup")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])

    add_body(doc, "This page illustrates the printable output generated from the Summary button.", size=9, color=COLORS["muted"])

    add_key_value_table(
        doc,
        [
            ["Booking Ref", "Q2026DEC01234", "Status", "Confirmed"],
            ["Quotation Ref", "Q2026DEC01234", "Agent", "Niran 1116"],
            ["Client Name", "Sample Client", "Pax", "2"],
            ["Booking Date", "14 Jul 2026", "Trip Period", "22 Dec 2026 - 29 Dec 2026"],
            ["Telephone", "0925991994", "Generated By", "Admin Office"],
        ],
    )

    add_service_table(
        doc,
        "Transfer Details",
        ["Date / Flight", "City", "Description", "From", "To", "TOT", "Status"],
        [
            ["22 Dec 2026", "Bangkok", "Airport to Hotel transfer", "Airport", "Hotel", "SIC", "Confirmed"],
            ["29 Dec 2026", "Bangkok", "Hotel to Airport transfer", "Hotel", "Airport", "SIC", "Confirmed"],
        ],
        [2.6, 2.2, 5.2, 2.2, 2.2, 1.5, 2.3],
    )

    add_service_table(
        doc,
        "Hotel Details",
        ["Check In", "Check Out", "City", "Hotel", "Room Type", "Nights", "Single", "Double", "Extra Bed", "Status"],
        [
            ["22 Dec 2026", "23 Dec 2026", "Bangkok", "Bizotel Premier - Bangkok", "Deluxe Room Incl. ABF (2A OC)", "1", "0", "1", "-", "Confirmed"],
            ["23 Dec 2026", "24 Dec 2026", "Bangkok", "Bizotel Premier - Bangkok", "Deluxe Room Incl. ABF (2A OC)", "1", "0", "1", "-", "Confirmed"],
        ],
        [2.0, 2.0, 1.8, 3.6, 3.3, 1.2, 1.2, 1.2, 1.5, 2.1],
    )

    add_service_table(
        doc,
        "Excursion Details",
        ["Date", "City", "Name", "Pickup", "Hotel", "TOE", "Status"],
        [
            ["24 Dec 2026", "Bangkok", "Railway Market & Floating Market", "08:00", "Bizotel Premier", "SIC", "Confirmed"],
        ],
        [2.3, 2.0, 5.3, 2.0, 3.4, 1.5, 2.3],
    )

    add_heading(doc, "Office Notes", 2)
    add_body(
        doc,
        "This summary is intended for internal office confirmation and supplier coordination after the booking has been confirmed. "
        "A customer-facing version can be generated later if needed by hiding internal notes and supplier-specific fields.",
        size=8.8,
    )

    doc.add_page_break()

    add_heading(doc, "4. Implementation Checklist", 1)
    checklist = [
        ("Add Summary button", "Place in Admin Edit Booking bottom action area, beside Booking Confirmed."),
        ("Confirm status gate", "Enable after booking is confirmed or when all services are confirmed."),
        ("Build print data source", "Use current booking payload including service tabs and confirmed flags."),
        ("Create A4 print view", "Hide all action buttons and use clean section tables."),
        ("Export/print", "Use browser print, PDF generation, or server-side document output."),
        ("Validation", "Compare final summary data with booking details before release."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_col_widths(table, [1.2, 5.0, 11.8])
    for i, h in enumerate(["No.", "Item", "Requirement"]):
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF", size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), COLORS["navy"])
    for idx, (item, requirement) in enumerate(checklist, 1):
        cells = table.add_row().cells
        set_cell_text(cells[0], idx, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], item, bold=True, size=8.5)
        set_cell_text(cells[2], requirement, size=8.5)
    style_table(table, header=False)

    add_heading(doc, "5. Acceptance Criteria", 1)
    add_bullet(doc, "Admin can click Summary after Booking Confirmed is available.")
    add_bullet(doc, "Generated A4 summary matches the booking detail page data.")
    add_bullet(doc, "All service sections print without horizontal overflow.")
    add_bullet(doc, "Action buttons and editable inputs are not shown in the printout.")
    add_bullet(doc, "Confirmed service rows are marked clearly, without using overly bright colors.")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
