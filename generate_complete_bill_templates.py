import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

DOCX_OUTPUT = "/Users/surachartlimrattanaphun/Desktop/tour/Bill_Template_VeraThailandia.docx"
XLSX_OUTPUT = "/Users/surachartlimrattanaphun/Desktop/tour/Bill_Template_VeraThailandia.xlsx"

FONT_NAME = "Tahoma"
NAVY = "1D3557"
ORANGE = "F4A261"
LIGHT = "F4F6F8"
MID = "D9E1E8"
WHITE = "FFFFFF"
TEXT = "243447"

def create_docx_template():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.35)
    section.right_margin = Cm(1.35)

    def set_cell_shading(cell, fill):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    def set_cell_border(cell, **edges):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge, attrs in edges.items():
            tag = "w:" + edge
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in attrs.items():
                element.set(qn("w:" + key), str(value))

    def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = tc_mar.find(qn("w:" + margin))
            if node is None:
                node = OxmlElement("w:" + margin)
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    def set_fixed_width(cell, width_cm):
        cell.width = Cm(width_cm)
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
        tc_w.set(qn("w:type"), "dxa")

    def set_font(run, size=9.5, bold=False, color=TEXT, italic=False):
        run.font.name = FONT_NAME
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor.from_string(color)

    def set_table_layout_fixed(table):
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.find(qn("w:tblLayout"))
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")

    # Header Title
    title_table = doc.add_table(rows=1, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(title_table)
    set_fixed_width(title_table.cell(0, 0), 12.0)
    set_fixed_width(title_table.cell(0, 1), 6.3)

    left = title_table.cell(0, 0)
    set_cell_shading(left, NAVY)
    set_cell_margins(left, top=140, bottom=140, start=160, end=160)
    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p_left.add_run("บิลเงินสด / ใบแจ้งหนี้" + chr(10))
    set_font(r1, size=20, bold=True, color=WHITE)
    r2 = p_left.add_run("CASH BILL / INVOICE")
    set_font(r2, size=11, bold=True, color=ORANGE)

    right = title_table.cell(0, 1)
    set_cell_shading(right, LIGHT)
    set_cell_margins(right, top=140, bottom=140, start=160, end=160)
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_no_lbl = p_right.add_run("เลขที่ / No. : ")
    set_font(r_no_lbl, size=9, bold=True, color=NAVY)
    r_no_val = p_right.add_run("[ระบุเลขที่บิล เช่น CB-202608-001]" + chr(10))
    set_font(r_no_val, size=9, color=TEXT, italic=True)

    r_dt_lbl = p_right.add_run("วันที่ / Date : ")
    set_font(r_dt_lbl, size=9, bold=True, color=NAVY)
    r_dt_val = p_right.add_run("[ระบุวันที่ เช่น 13 สิงหาคม 2026]")
    set_font(r_dt_val, size=9, color=TEXT, italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Info Table (Seller & Buyer Headers)
    info = doc.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(info)
    for row in info.rows:
        set_fixed_width(row.cells[0], 9.15)
        set_fixed_width(row.cells[1], 9.15)
        for cell in row.cells:
            set_cell_border(cell, top={"val": "single", "sz": 4, "color": MID},
                            bottom={"val": "single", "sz": 4, "color": MID},
                            start={"val": "single", "sz": 4, "color": MID},
                            end={"val": "single", "sz": 4, "color": MID})
            set_cell_margins(cell, top=80, bottom=80, start=120, end=120)

    # Column Headers for Seller/Buyer
    p_s_hdr = info.cell(0, 0).paragraphs[0]
    r = p_s_hdr.add_run("ข้อมูลผู้ขาย / ผู้รับเงิน (Seller / Billed From)")
    set_font(r, size=9.5, bold=True, color=WHITE)
    set_cell_shading(info.cell(0, 0), NAVY)

    p_b_hdr = info.cell(0, 1).paragraphs[0]
    r = p_b_hdr.add_run("ข้อมูลผู้ซื้อ / ลูกค้า (Customer / Billed To)")
    set_font(r, size=9.5, bold=True, color=WHITE)
    set_cell_shading(info.cell(0, 1), NAVY)

    # Seller Rows
    p = info.cell(1, 0).paragraphs[0]
    p.add_run("ชื่อ / Name: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("คุณสุรชาติ ลิมป์รัตนพุน (Surachart Limrattanaphun)")

    p = info.cell(2, 0).paragraphs[0]
    p.add_run("ที่อยู่ / Address: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("[ระบุที่อยู่ผู้รับเงิน/ผู้ขาย]")

    p = info.cell(3, 0).paragraphs[0]
    p.add_run("เลขผู้เสียภาษี/บัตรประชาชน: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("[ระบุเลขประจำตัวผู้เสียภาษีหรือเลขบัตรประชาชน]")

    p = info.cell(4, 0).paragraphs[0]
    p.add_run("โทรศัพท์ / Phone: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("[ระบุเบอร์โทรศัพท์]")

    # Buyer Rows (Pre-filled VeraThailandia)
    p = info.cell(1, 1).paragraphs[0]
    p.add_run("ชื่อบริษัท / Name: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("บริษัท เวร่าไทยแลนด์เดีย จำกัด (VeraThailandia Co., Ltd.)")

    p = info.cell(2, 1).paragraphs[0]
    p.add_run("ที่อยู่ / Address: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("160/424-425 อาคารไอทีเอฟ สีลมพาเลส ชั้น 20 ถนนสีลม แขวงสุริยวงศ์ เขตบางรัก กรุงเทพฯ 10500")

    p = info.cell(3, 1).paragraphs[0]
    p.add_run("เลขประจำตัวผู้เสียภาษี / Tax ID: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("0105547045569")

    p = info.cell(4, 1).paragraphs[0]
    p.add_run("โทรศัพท์ / อ้างอิง: ").bold = True
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.add_run("(+66 2) 126 6914 | โครงการ Tour Management System")

    # Section Heading
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(8)
    p_sec.paragraph_format.space_after = Pt(4)
    r = p_sec.add_run("รายการ / DESCRIPTION")
    set_font(r, size=10.5, bold=True, color=NAVY)

    # Items Table
    items_table = doc.add_table(rows=6, cols=5)
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(items_table)
    widths = [1.2, 9.1, 1.8, 2.8, 3.4]
    headers = ["ลำดับ" + chr(10) + "No.", "รายละเอียด" + chr(10) + "Description", "จำนวน" + chr(10) + "Qty", "ราคาต่อหน่วย" + chr(10) + "Unit Price", "จำนวนเงิน" + chr(10) + "Amount (THB)"]
    
    for row in items_table.rows:
        for idx, w in enumerate(widths):
            set_fixed_width(row.cells[idx], w)

    # Header Row
    for idx, cell in enumerate(items_table.rows[0].cells):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=100, bottom=100, start=60, end=60)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headers[idx])
        set_font(r, size=8.5, bold=True, color=WHITE)

    # Sample rows
    sample_data = [
        ("1", "ค่าพัฒนาระบบ Tour Management System (งวดที่ ...)", "1", "[ระบุจำนวนเงิน]", "[ระบุจำนวนเงิน]"),
        ("2", "[รายละเอียดบริการเพิ่มเติม เช่น ค่าติดตั้ง/ค่าดูแลระบบ]", "-", "-", "-"),
        ("3", "", "", "", ""),
        ("4", "", "", "", ""),
        ("5", "", "", "", "")
    ]

    for r_idx, data in enumerate(sample_data, start=1):
        row_cells = items_table.rows[r_idx].cells
        bg_color = WHITE if r_idx % 2 != 0 else "FAFBFC"
        for c_idx, text in enumerate(data):
            cell = row_cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_margins(cell, top=90, bottom=90, start=80, end=80)
            set_cell_border(cell, top={"val": "single", "sz": 4, "color": MID},
                            bottom={"val": "single", "sz": 4, "color": MID},
                            start={"val": "single", "sz": 4, "color": MID},
                            end={"val": "single", "sz": 4, "color": MID})
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 2) else (WD_ALIGN_PARAGRAPH.RIGHT if c_idx in (3, 4) else WD_ALIGN_PARAGRAPH.LEFT)
            p.alignment = align
            r = p.add_run(text)
            set_font(r, size=9, color=TEXT if not text.startswith("[") else "888888")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Totals Table
    totals_table = doc.add_table(rows=4, cols=2)
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_layout_fixed(totals_table)
    set_fixed_width(totals_table.cell(0, 0), 14.9)
    set_fixed_width(totals_table.cell(0, 1), 3.4)

    tot_labels = [
        ("รวมเป็นเงิน / Subtotal", "[ระบุจำนวนเงิน]"),
        ("ส่วนลด / Discount", "-"),
        ("ภาษีมูลค่าเพิ่ม / VAT (ถ้ามี)", "-"),
        ("ยอดรวมสุทธิ / TOTAL AMOUNT", "[ระบุจำนวนเงินสุทธิ]")
    ]

    for i, (label, val) in enumerate(tot_labels):
        c0 = totals_table.cell(i, 0)
        c1 = totals_table.cell(i, 1)
        fill = NAVY if i == 3 else LIGHT
        set_cell_shading(c0, fill)
        set_cell_shading(c1, fill)
        set_cell_margins(c0, top=80, bottom=80, start=100, end=100)
        set_cell_margins(c1, top=80, bottom=80, start=100, end=100)
        
        for c in (c0, c1):
            set_cell_border(c, top={"val": "single", "sz": 4, "color": MID},
                            bottom={"val": "single", "sz": 4, "color": MID},
                            start={"val": "single", "sz": 4, "color": MID},
                            end={"val": "single", "sz": 4, "color": MID})

        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label)
        set_font(r0, size=9, bold=(i == 3), color=WHITE if i == 3 else NAVY)

        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = p1.add_run(val)
        set_font(r1, size=9.5, bold=(i == 3), color=WHITE if i == 3 else TEXT)

    # Amount in words
    words_table = doc.add_table(rows=1, cols=1)
    words_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(words_table)
    set_fixed_width(words_table.cell(0, 0), 18.3)
    c_w = words_table.cell(0, 0)
    set_cell_shading(c_w, "FFF6ED")
    set_cell_margins(c_w, top=100, bottom=100, start=120, end=120)
    set_cell_border(c_w, top={"val": "single", "sz": 6, "color": ORANGE},
                    bottom={"val": "single", "sz": 6, "color": ORANGE},
                    start={"val": "single", "sz": 6, "color": ORANGE},
                    end={"val": "single", "sz": 6, "color": ORANGE})
    p_w = c_w.paragraphs[0]
    r_w_lbl = p_w.add_run("จำนวนเงิน (ตัวอักษร) / Amount in words: ")
    set_font(r_w_lbl, size=9, bold=True, color=NAVY)
    r_w_val = p_w.add_run("[ระบุจำนวนเงินเป็นตัวอักษรภาษาไทย เช่น (หนึ่งแสนบาทถ้วน)]")
    set_font(r_w_val, size=9.5, color=TEXT, italic=True)

    # Payment details & Notes
    p_pay = doc.add_paragraph()
    p_pay.paragraph_format.space_before = Pt(8)
    p_pay.paragraph_format.space_after = Pt(2)
    r = p_pay.add_run("วิธีการชำระเงิน / Payment Method:" + chr(10))
    set_font(r, size=9.2, bold=True, color=NAVY)
    r2 = p_pay.add_run("☐ เงินสด / Cash     ☑ โอนเงินผ่านธนาคาร / Bank Transfer" + chr(10))
    set_font(r2, size=9)
    r3 = p_pay.add_run("ข้อมูลบัญชีธนาคาร: ")
    set_font(r3, size=9, bold=True, color=NAVY)
    r4 = p_pay.add_run("ธนาคาร [ระบุชื่อธนาคาร] | เลขที่บัญชี [ระบุเลขบัญชี] | ชื่อบัญชี [สุรชาติ ลิมป์รัตนพุน]")
    set_font(r4, size=9)

    # Signatures
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(sig_table)
    for row in sig_table.rows:
        for cell in row.cells:
            set_fixed_width(cell, 6.1)
            set_cell_margins(cell, top=100, bottom=100)

    sig_labels = ["ผู้จ่ายเงิน / Payer", "ผู้รับเงิน / Receiver", "ผู้มีอำนาจลงนาม / Authorized"]
    for i, label in enumerate(sig_labels):
        p0 = sig_table.cell(0, i).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p0.add_run("ลงชื่อ / Signature" + chr(10) + chr(10) + "...........................................")
        set_font(r, size=8.5, color="777777")

        p1 = sig_table.cell(1, i).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(label + chr(10) + "วันที่ / Date: ........................")
        set_font(r, size=8.5, bold=True, color=NAVY)

    doc.save(DOCX_OUTPUT)
    print("DOCX created:", DOCX_OUTPUT)


def create_xlsx_template():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cash_Bill_Template"
    ws.views.sheetView[0].showGridLines = True

    # Styling
    navy_fill = PatternFill(start_color="1D3557", end_color="1D3557", fill_type="solid")
    light_fill = PatternFill(start_color="F4F6F8", end_color="F4F6F8", fill_type="solid")
    orange_fill = PatternFill(start_color="FFF6ED", end_color="FFF6ED", fill_type="solid")
    
    font_hdr = Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")
    font_bold_navy = Font(name=FONT_NAME, size=9.5, bold=True, color="1D3557")
    font_regular = Font(name=FONT_NAME, size=9.5, color="243447")
    font_placeholder = Font(name=FONT_NAME, size=9.5, italic=True, color="777777")

    thin_border = Border(
        left=Side(style='thin', color='D9E1E8'),
        right=Side(style='thin', color='D9E1E8'),
        top=Side(style='thin', color='D9E1E8'),
        bottom=Side(style='thin', color='D9E1E8')
    )

    # Set Column Widths
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 42
    ws.column_dimensions['C'].width = 10
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 18

    # Title Block
    ws.merge_cells('A1:C2')
    title_cell = ws['A1']
    title_cell.value = "บิลเงินสด / ใบแจ้งหนี้ (CASH BILL / INVOICE)"
    title_cell.font = Font(name=FONT_NAME, size=13, bold=True, color="FFFFFF")
    title_cell.fill = navy_fill
    title_cell.alignment = Alignment(horizontal='center', vertical='center')

    for r in range(1, 3):
        for c in range(1, 4):
            ws.cell(row=r, column=c).fill = navy_fill

    ws.merge_cells('D1:E1')
    ws['D1'] = "เลขที่ / No.: [CB-202608-001]"
    ws['D1'].font = font_bold_navy
    ws['D1'].fill = light_fill
    ws['D1'].alignment = Alignment(horizontal='left', vertical='center')

    ws.merge_cells('D2:E2')
    ws['D2'] = "วันที่ / Date: [13 สิงหาคม 2026]"
    ws['D2'].font = font_bold_navy
    ws['D2'].fill = light_fill
    ws['D2'].alignment = Alignment(horizontal='left', vertical='center')

    ws.row_dimensions[1].height = 22
    ws.row_dimensions[2].height = 22

    # Seller & Buyer Section
    ws.merge_cells('A4:B4')
    ws['A4'] = "ข้อมูลผู้ขาย / ผู้รับเงิน (Seller / Billed From)"
    ws['A4'].font = font_hdr
    ws['A4'].fill = navy_fill

    ws.merge_cells('C4:E4')
    ws['C4'] = "ข้อมูลผู้ซื้อ / ลูกค้า (Customer / Billed To)"
    ws['C4'].font = font_hdr
    ws['C4'].fill = navy_fill

    seller_buyer_data = [
        ("ชื่อ / Name:", "คุณสุรชาติ ลิมป์รัตนพุน", "ชื่อบริษัท:", "บริษัท เวร่าไทยแลนด์เดีย จำกัด (VeraThailandia Co., Ltd.)"),
        ("ที่อยู่ / Address:", "[ระบุที่อยู่ผู้รับเงิน/ผู้ขาย]", "ที่อยู่ / Address:", "160/424-425 อาคารไอทีเอฟ สีลมพาเลส ชั้น 20 ถ.สีลม เขตบางรัก กทม."),
        ("เลขผู้เสียภาษี:", "[ระบุเลขประจำตัวผู้เสียภาษี/บัตรประชาชน]", "เลขผู้เสียภาษี:", "0105547045569"),
        ("โทรศัพท์:", "[ระบุเบอร์โทรศัพท์]", "โทรศัพท์/อ้างอิง:", "(+66 2) 126 6914 | โครงการ Tour Management System")
    ]

    for idx, (s_lbl, s_val, b_lbl, b_val) in enumerate(seller_buyer_data, start=5):
        ws.cell(row=idx, column=1, value=s_lbl).font = font_bold_navy
        ws.cell(row=idx, column=2, value=s_val).font = font_regular if not s_val.startswith("[") else font_placeholder
        
        ws.cell(row=idx, column=3, value=b_lbl).font = font_bold_navy
        ws.merge_cells(start_row=idx, start_column=4, end_row=idx, end_column=5)
        ws.cell(row=idx, column=4, value=b_val).font = font_regular

        for c in range(1, 6):
            ws.cell(row=idx, column=c).border = thin_border

    # Item Table Headers
    ws.row_dimensions[10].height = 25
    headers = [("A10", "ลำดับ" + chr(10) + "No."), ("B10", "รายละเอียด" + chr(10) + "Description"), ("C10", "จำนวน" + chr(10) + "Qty"), ("D10", "ราคา/หน่วย" + chr(10) + "Unit Price"), ("E10", "จำนวนเงิน" + chr(10) + "Amount (THB)")]
    for cell_pos, text in headers:
        cell = ws[cell_pos]
        cell.value = text
        cell.font = font_hdr
        cell.fill = navy_fill
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # Sample Data Rows
    items = [
        (1, "ค่าพัฒนาระบบ Tour Management System (งวดที่ 1)", 1, 50000),
        (2, "[ระบุรายละเอียดบริการเพิ่มเติม]", 1, 0),
        (3, "", None, None),
        (4, "", None, None),
        (5, "", None, None)
    ]

    for idx, (no, desc, qty, price) in enumerate(items, start=11):
        ws.row_dimensions[idx].height = 20
        ws.cell(row=idx, column=1, value=no if desc else "").alignment = Alignment(horizontal='center')
        ws.cell(row=idx, column=2, value=desc).alignment = Alignment(horizontal='left')
        ws.cell(row=idx, column=3, value=qty if qty is not None else "").alignment = Alignment(horizontal='center')
        ws.cell(row=idx, column=4, value=price if price is not None else "").alignment = Alignment(horizontal='right')
        
        # Formula for Amount
        if qty is not None and price is not None:
            ws.cell(row=idx, column=5, value="=C" + str(idx) + "*D" + str(idx)).alignment = Alignment(horizontal='right')
            ws.cell(row=idx, column=5).number_format = '#,##0.00'
            ws.cell(row=idx, column=4).number_format = '#,##0.00'

        for c in range(1, 6):
            cell = ws.cell(row=idx, column=c)
            cell.border = thin_border
            cell.font = font_regular if not str(cell.value).startswith("[") else font_placeholder

    # Totals
    ws.merge_cells('A16:D16')
    ws['A16'] = "รวมเป็นเงิน / Subtotal"
    ws['A16'].alignment = Alignment(horizontal='right')
    ws['A16'].font = font_bold_navy
    ws['E16'] = "=SUM(E11:E15)"
    ws['E16'].font = font_bold_navy
    ws['E16'].number_format = '#,##0.00'

    ws.merge_cells('A17:D17')
    ws['A17'] = "ภาษีมูลค่าเพิ่ม / VAT (7%)"
    ws['A17'].alignment = Alignment(horizontal='right')
    ws['A17'].font = font_bold_navy
    ws['E17'] = 0
    ws['E17'].number_format = '#,##0.00'

    ws.merge_cells('A18:D18')
    ws['A18'] = "ยอดรวมสุทธิ / TOTAL AMOUNT"
    ws['A18'].alignment = Alignment(horizontal='right')
    ws['A18'].font = Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")
    ws['A18'].fill = navy_fill
    ws['E18'] = "=E16+E17"
    ws['E18'].font = Font(name=FONT_NAME, size=10, bold=True, color="FFFFFF")
    ws['E18'].fill = navy_fill
    ws['E18'].number_format = '#,##0.00'

    for r in range(16, 19):
        for c in range(1, 6):
            ws.cell(row=r, column=c).border = thin_border

    # Amount in words
    ws.merge_cells('A20:E20')
    ws['A20'] = "จำนวนเงิน (ตัวอักษร) / Amount in words: [ระบุจำนวนเงินเป็นตัวอักษร เช่น (ห้าหมื่นบาทถ้วน)]"
    ws['A20'].font = font_placeholder
    ws['A20'].fill = orange_fill
    ws['A20'].alignment = Alignment(horizontal='left', vertical='center')

    # Payment details
    ws.cell(row=22, column=1, value="วิธีการชำระเงิน:").font = font_bold_navy
    ws.cell(row=22, column=2, value="[X] โอนเงินผ่านธนาคาร  [ ] เงินสด").font = font_regular
    ws.cell(row=23, column=1, value="ข้อมูลบัญชี:").font = font_bold_navy
    ws.cell(row=23, column=2, value="ธนาคาร [ชื่อธนาคาร] | เลขบัญชี [เลขบัญชี] | ชื่อบัญชี สุรชาติ ลิมป์รัตนพุน").font = font_regular

    wb.save(XLSX_OUTPUT)
    print("XLSX created:", XLSX_OUTPUT)

if __name__ == "__main__":
    create_docx_template()
    create_xlsx_template()
