from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "/Users/surachartlimrattanaphun/Desktop/tour/Cash_Bill_Template_Verathailandia.docx"
FONT = "Arial Unicode MS"
NAVY = "1D3557"
ORANGE = "F4A261"
LIGHT = "F4F6F8"
MID = "D9E1E8"
WHITE = "FFFFFF"
TEXT = "243447"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in edges.items():
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in attrs.items():
            element.set(qn("w:" + key), str(value))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + margin))
        if node is None:
            node = OxmlElement("w:" + margin)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_fixed_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=10, bold=False, color=TEXT, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.0):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    style_paragraph(p)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return p


def add_cell_text(cell, text, *, size=9.5, bold=False, color=TEXT,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = clear_cell(cell)
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def set_table_layout_fixed(table):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_labeled_line(cell, label, placeholder=""):
    p = clear_cell(cell)
    r = p.add_run(label)
    set_font(r, size=9, bold=True, color=NAVY)
    r = p.add_run(placeholder or "  ")
    set_font(r, size=9.5, color=TEXT)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.25)
section.bottom_margin = Cm(1.25)
section.left_margin = Cm(1.35)
section.right_margin = Cm(1.35)
section.header_distance = Cm(0.55)
section.footer_distance = Cm(0.55)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(0)

# Quiet running header.
header_p = section.header.paragraphs[0]
style_paragraph(header_p, WD_ALIGN_PARAGRAPH.RIGHT)
set_font(header_p.add_run("CASH BILL TEMPLATE"), size=8, bold=True, color="7A8793")

# Title block.
title_table = doc.add_table(rows=1, cols=2)
title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_layout_fixed(title_table)
set_fixed_width(title_table.cell(0, 0), 12.0)
set_fixed_width(title_table.cell(0, 1), 6.0)

left = title_table.cell(0, 0)
set_cell_shading(left, NAVY)
left_p = clear_cell(left)
left_p.paragraph_format.space_before = Pt(5)
left_p.paragraph_format.space_after = Pt(5)
set_font(left_p.add_run("บิลเงินสด"), size=23, bold=True, color=WHITE)
left_p.add_run("\n")
set_font(left_p.add_run("CASH BILL"), size=11, bold=True, color=ORANGE)

right = title_table.cell(0, 1)
set_cell_shading(right, LIGHT)
rp = clear_cell(right)
rp.paragraph_format.space_before = Pt(2)
rp.paragraph_format.space_after = Pt(2)
set_font(rp.add_run("เลขที่ / No.\n"), size=8.5, bold=True, color=NAVY)
set_font(rp.add_run("................................\n"), size=10, color=TEXT)
set_font(rp.add_run("วันที่ / Date\n"), size=8.5, bold=True, color=NAVY)
set_font(rp.add_run("................................"), size=10, color=TEXT)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Seller / buyer data.
info = doc.add_table(rows=4, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_layout_fixed(info)
for row in info.rows:
    set_fixed_width(row.cells[0], 9.0)
    set_fixed_width(row.cells[1], 9.0)
    for cell in row.cells:
        set_cell_border(cell, top={"val": "single", "sz": 6, "color": MID},
                        bottom={"val": "single", "sz": 6, "color": MID},
                        start={"val": "single", "sz": 6, "color": MID},
                        end={"val": "single", "sz": 6, "color": MID})

add_labeled_line(info.cell(0, 0), "ผู้ขาย / Seller: ", "........................................................")
add_labeled_line(info.cell(0, 1), "ผู้ซื้อ / Customer: ", "....................................................")
add_labeled_line(info.cell(1, 0), "ที่อยู่ / Address: ", ".......................................................")
add_labeled_line(info.cell(1, 1), "ที่อยู่ / Address: ", ".......................................................")
add_labeled_line(info.cell(2, 0), "เลขประจำตัวผู้เสียภาษี / Tax ID: ", "................................")
add_labeled_line(info.cell(2, 1), "เลขประจำตัวผู้เสียภาษี / Tax ID: ", "................................")
add_labeled_line(info.cell(3, 0), "โทรศัพท์ / Telephone: ", "...............................................")
add_labeled_line(info.cell(3, 1), "อ้างอิง / Reference: ", "....................................................")

section_heading = doc.add_paragraph()
style_paragraph(section_heading, before=8, after=4)
set_font(section_heading.add_run("รายการ / DESCRIPTION"), size=10.5, bold=True, color=NAVY)

# Main item table.
table = doc.add_table(rows=8, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_layout_fixed(table)
widths = [1.2, 9.0, 2.0, 2.8, 3.0]
headers = ["ลำดับ\nNo.", "รายละเอียด\nDescription", "จำนวน\nQty", "ราคาต่อหน่วย\nUnit Price", "จำนวนเงิน\nAmount"]
for col, width in enumerate(widths):
    for row in table.rows:
        set_fixed_width(row.cells[col], width)

for idx, cell in enumerate(table.rows[0].cells):
    set_cell_shading(cell, NAVY)
    add_cell_text(cell, headers[idx], size=8.5, bold=True, color=WHITE,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_margins(cell, top=120, bottom=120)
set_repeat_table_header(table.rows[0])

for r in range(1, 8):
    for c, cell in enumerate(table.rows[r].cells):
        set_cell_shading(cell, WHITE if r % 2 else "FAFBFC")
        add_cell_text(cell, str(r) if c == 0 else "", size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER if c != 1 else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_margins(cell, top=125, bottom=125)
        set_cell_border(cell, top={"val": "single", "sz": 6, "color": MID},
                        bottom={"val": "single", "sz": 6, "color": MID},
                        start={"val": "single", "sz": 6, "color": MID},
                        end={"val": "single", "sz": 6, "color": MID})

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Totals and amount in words.
totals = doc.add_table(rows=4, cols=2)
totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
set_table_layout_fixed(totals)
for row in totals.rows:
    set_fixed_width(row.cells[0], 14.9)
    set_fixed_width(row.cells[1], 3.1)

labels = ["รวมเป็นเงิน / Subtotal", "ส่วนลด / Discount", "ภาษีมูลค่าเพิ่ม / VAT", "ยอดรวมสุทธิ / TOTAL"]
for i, label in enumerate(labels):
    add_cell_text(totals.cell(i, 0), label, size=9.2, bold=(i == 3),
                  align=WD_ALIGN_PARAGRAPH.RIGHT, color=WHITE if i == 3 else NAVY)
    add_cell_text(totals.cell(i, 1), "", size=9.5, bold=(i == 3),
                  align=WD_ALIGN_PARAGRAPH.RIGHT, color=WHITE if i == 3 else TEXT)
    fill = NAVY if i == 3 else LIGHT
    set_cell_shading(totals.cell(i, 0), fill)
    set_cell_shading(totals.cell(i, 1), fill)
    for cell in totals.rows[i].cells:
        set_cell_border(cell, top={"val": "single", "sz": 6, "color": MID},
                        bottom={"val": "single", "sz": 6, "color": MID},
                        start={"val": "single", "sz": 6, "color": MID},
                        end={"val": "single", "sz": 6, "color": MID})

words = doc.add_table(rows=1, cols=1)
words.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_layout_fixed(words)
set_fixed_width(words.cell(0, 0), 18.0)
add_labeled_line(words.cell(0, 0), "จำนวนเงิน (ตัวอักษร) / Amount in words: ", ".................................................................")
set_cell_shading(words.cell(0, 0), "FFF6ED")
set_cell_border(words.cell(0, 0), top={"val": "single", "sz": 8, "color": ORANGE},
                bottom={"val": "single", "sz": 8, "color": ORANGE},
                start={"val": "single", "sz": 8, "color": ORANGE},
                end={"val": "single", "sz": 8, "color": ORANGE})

payment = doc.add_paragraph()
style_paragraph(payment, before=7, after=1)
set_font(payment.add_run("วิธีชำระเงิน / Payment Method:  "), size=9.2, bold=True, color=NAVY)
set_font(payment.add_run("☐ เงินสด / Cash     ☐ โอนเงิน / Transfer     ☐ อื่น ๆ / Other: ........................"), size=9.2)

note = doc.add_paragraph()
style_paragraph(note, after=5)
set_font(note.add_run("หมายเหตุ / Note: "), size=9.2, bold=True, color=NAVY)
set_font(note.add_run("............................................................................................................................"), size=9.2)

# Signatures.
sig = doc.add_table(rows=2, cols=3)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_layout_fixed(sig)
for row in sig.rows:
    for cell in row.cells:
        set_fixed_width(cell, 6.0)
        set_cell_margins(cell, top=100, bottom=100)

for i, label in enumerate(("ผู้จ่ายเงิน / Payer", "ผู้รับเงิน / Receiver", "ผู้มีอำนาจลงนาม / Authorized")):
    add_cell_text(sig.cell(0, i), "ลงชื่อ / Signature\n\n........................................", size=8.8,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(sig.cell(1, i), label + "\nวันที่ / Date: ........................", size=8.5,
                  bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

footer = section.footer.paragraphs[0]
style_paragraph(footer, WD_ALIGN_PARAGRAPH.CENTER)
set_font(footer.add_run("เอกสารฉบับนี้เป็นแบบฟอร์มเปล่าสำหรับกรอกข้อมูล / Editable blank template"),
         size=7.5, color="7A8793")

# Prevent accidental second blank page from trailing paragraph spacing.
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.keep_together = True

doc.save(OUTPUT)
print(OUTPUT)
